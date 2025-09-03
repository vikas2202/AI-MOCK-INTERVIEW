# import streamlit as st
# import speech_recognition as sr
# import PyPDF2
# import docx
# import time
# from openai import OpenAI
# from openai import OpenAIError

# # ===============================
# # CONFIG
# # ===============================
# st.set_page_config(page_title="AI Mock Interviewer", page_icon="🧠", layout="wide")

# # Replace with your provided API key
# client = OpenAI(api_key="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")

# # ===============================
# # FUNCTIONS
# # ===============================

# # Extract text from resume
# def extract_text(file):
#     text = ""
#     if file.type == "application/pdf":
#         reader = PyPDF2.PdfReader(file)
#         for page in reader.pages:
#             text += page.extract_text() or ""
#     elif file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
#         doc = docx.Document(file)
#         for para in doc.paragraphs:
#             text += para.text + "\n"
#     else:
#         text = file.read().decode("utf-8")
#     return text.strip()

# # Handle OpenAI API call with retry for rate limits
# def call_openai(func, *args, retries=3, delay=5, **kwargs):
#     for attempt in range(retries):
#         try:
#             return func(*args, **kwargs)
#         except OpenAIError as e:
#             if "RateLimit" in str(e):
#                 if attempt < retries - 1:
#                     st.warning(f"⚠️ Rate limit reached. Retrying in {delay} seconds...")
#                     time.sleep(delay)
#                 else:
#                     st.error("❌ API rate limit exceeded. Try again later.")
#                     return None
#             else:
#                 st.error(f"❌ OpenAI API Error: {e}")
#                 return None

# # AI generated random question
# def get_ai_question(resume_text, asked_qs):
#     # First question: always introduction
#     if not asked_qs:
#         return "Please introduce yourself."

#     prompt = f"""
# You are an AI interviewer. Based on this resume: {resume_text},
# generate **one random interview question** (technical or behavioral or basic coding question).
# Do NOT repeat previous questions. Previous asked questions: {asked_qs}
# Level: start from basic, then intermediate, then advanced
# """
#     response = call_openai(
#         client.chat.completions.create,
#         model="gpt-4o-mini",
#         messages=[{"role": "user", "content": prompt}]
#     )
#     if response:
#         return response.choices[0].message.content.strip()
#     return "❌ Could not generate question due to API issue."

# # AI feedback generator
# def get_feedback(questions_answers):
#     prompt = f"""
# I conducted a mock interview. Here are the questions and answers:
# {questions_answers}

# Give feedback in 2 clear sections:
# ✅ Strengths
# ❌ Improvements
# """
#     response = call_openai(
#         client.chat.completions.create,
#         model="gpt-4o-mini",
#         messages=[{"role": "user", "content": prompt}]
#     )
#     if response:
#         return response.choices[0].message.content.strip()
#     return "❌ Could not generate feedback due to API issue."

# # Record speech and convert to text
# def record_answer(min_seconds=8):
#     recognizer = sr.Recognizer()
#     with sr.Microphone() as source:
#         st.info("🎤 Listening... Please speak now (minimum 8 seconds)")
#         audio = recognizer.listen(source)
#         try:
#             text = recognizer.recognize_google(audio)
#             duration = len(text.split()) / 2  # rough estimate: 2 words/sec
#             if duration < min_seconds:
#                 st.warning(f"⚠️ Your answer is too short. Try to speak at least {min_seconds} seconds.")
#             st.success(f"📝 Recognized: {text}")
#             return text
#         except:
#             st.error("⚠️ Could not recognize your speech. Try again.")
#             return None

# # ===============================
# # STREAMLIT APP
# # ===============================
# st.title("🧠 AI Mock Interviewer with Speech & Text Input")

# # Session state
# if "resume_text" not in st.session_state:
#     st.session_state.resume_text = ""
# if "questions" not in st.session_state:
#     st.session_state.questions = []
# if "answers" not in st.session_state:
#     st.session_state.answers = []
# if "interview_finished" not in st.session_state:
#     st.session_state.interview_finished = False

# # Upload Resume
# uploaded_file = st.file_uploader("📂 Upload your Resume (PDF/DOCX/TXT)", type=["pdf", "docx", "txt"])
# if uploaded_file and not st.session_state.resume_text:
#     st.session_state.resume_text = extract_text(uploaded_file)
#     st.success("✅ Resume uploaded & processed!")

# # Start Interview
# if st.session_state.resume_text and not st.session_state.interview_finished:

#     if st.button("🤖 Ask Next Question"):
#         q = get_ai_question(st.session_state.resume_text, st.session_state.questions)
#         st.session_state.questions.append(q)
#         st.session_state.answers.append("")  # placeholder

#     # Display current question and record/write answer
#     if st.session_state.questions:
#         q_index = len(st.session_state.questions) - 1
#         st.subheader(f"Q{q_index+1}: {st.session_state.questions[q_index]}")

#         # Option to write or speak answer
#         answer_choice = st.radio("Choose your answer method:", ("Type Answer", "Speak Answer"))

#         if answer_choice == "Type Answer":
#             typed_answer_key = f"typed_answer_{q_index}"
#             if typed_answer_key not in st.session_state:
#                 st.session_state[typed_answer_key] = ""
#             typed_answer = st.text_area(
#                 "📝 Write your answer here:",
#                 value=st.session_state[typed_answer_key],
#                 height=150,
#                 key=typed_answer_key
#             )
#             if st.button("💾 Save Answer"):
#                 st.session_state.answers[q_index] = typed_answer
#                 st.success("✅ Answer saved!")
#                 # Clear text area after saving
#                 st.session_state[typed_answer_key] = ""

#         elif answer_choice == "Speak Answer":
#             if st.button("🎤 Record My Answer"):
#                 ans = record_answer()
#                 if ans:
#                     st.session_state.answers[q_index] = ans

#     # Show Q&A so far
#     if st.session_state.answers:
#         st.subheader("📌 Your Progress")
#         for i, (q, a) in enumerate(zip(st.session_state.questions, st.session_state.answers)):
#             st.markdown(f"**Q{i+1}:** {q}")
#             st.markdown(f"**A{i+1}:** {a if a else '❌ Not answered yet'}")

#     # Stop interview anytime
#     if st.button("🛑 Stop Interview & Get Feedback"):
#         st.session_state.interview_finished = True

# # Finish Interview automatically after 10 questions or stopped manually
# if st.session_state.interview_finished or len(st.session_state.questions) >= 10:
#     qa_text = "\n".join([f"Q: {q}\nA: {a}" for q, a in zip(st.session_state.questions, st.session_state.answers)])
#     feedback = get_feedback(qa_text)
#     st.subheader("📋 Interview Feedback")
#     st.write(feedback)















# import streamlit as st
# import speech_recognition as sr
# import PyPDF2
# import docx
# import time
# from openai import OpenAI
# from openai import OpenAIError

# # ===============================
# # CONFIG
# # ===============================
# st.set_page_config(page_title="AI Mock Interviewer", page_icon="🧠", layout="wide")

# # ===============================
# # CSS (Streamlit-compatible)
# # ===============================
# st.markdown("""
# <style>
# /* Background gradient for app */
# [data-testid="stAppViewContainer"] {
#     background: linear-gradient(135deg, #0f2027, #203a43, #2c5364);
#     color: white;
#     font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
# }

# /* Title */
# h1 {
#     text-align: center;
#     font-size: 3rem;
#     background: -webkit-linear-gradient(#00c6ff, #0072ff);
#     -webkit-background-clip: text;
#     -webkit-text-fill-color: transparent;
#     margin-bottom: 25px;
# }

# /* Buttons */
# div.stButton > button {
#     background: linear-gradient(45deg, #ff416c, #ff4b2b);
#     border-radius: 12px;
#     color: black;
#     font-size: 16px;
#     font-weight: bold;
#     padding: 0.6rem 1.2rem;
#     border: none;
#     transition: all 0.3s ease-in-out;
# }
# div.stButton > button:hover {
#     transform: scale(1.05);
#     background: linear-gradient(45deg, #ff4b2b, #ff416c);
# }

# /* Radio buttons */
# div[role="radiogroup"] label {
#     color: #f0f0f0;
#     font-weight: 600;
# }

# /* Text area */
# textarea {
#     background: #1e293b !important;
#     color: #f8f9fa !important;
#     border-radius: 12px !important;
#     padding: 10px !important;
#     font-size: 15px !important;
# }

# /* File uploader */
# div.stFileUploader label {
#     color: #00c6ff !important;
#     font-weight: 600;
#     font-size: 1.1rem;
# }

# /* Q&A card */
# .qa-card {
#     background: #1e1e2f;
#     border-radius: 12px;
#     padding: 15px;
#     margin-top: 15px;
#     color: #f8f9fa;
#     box-shadow: 0px 4px 15px rgba(0,0,0,0.5);
# }

# /* Feedback box */
# .feedback-box {
#     background: #1e1e2f;
#     border-radius: 12px;
#     padding: 20px;
#     margin-top: 15px;
#     color: #f8f9fa;
#     box-shadow: 0px 4px 15px rgba(0,0,0,0.5);
# }
# </style>
# """, unsafe_allow_html=True)

# # ===============================
# # OPENAI CONFIG
# # ===============================
# client = OpenAI(api_key="YOUR_OPENAI_API_KEY")  # Replace with your API key

# # ===============================
# # FUNCTIONS (same as your original)
# # ===============================
# def extract_text(file):
#     text = ""
#     if file.type == "application/pdf":
#         reader = PyPDF2.PdfReader(file)
#         for page in reader.pages:
#             text += page.extract_text() or ""
#     elif file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
#         doc = docx.Document(file)
#         for para in doc.paragraphs:
#             text += para.text + "\n"
#     else:
#         text = file.read().decode("utf-8")
#     return text.strip()

# def call_openai(func, *args, retries=3, delay=5, **kwargs):
#     for attempt in range(retries):
#         try:
#             return func(*args, **kwargs)
#         except OpenAIError as e:
#             if "RateLimit" in str(e):
#                 if attempt < retries - 1:
#                     st.warning(f"⚠️ Rate limit reached. Retrying in {delay} seconds...")
#                     time.sleep(delay)
#                 else:
#                     st.error("❌ API rate limit exceeded. Try again later.")
#                     return None
#             else:
#                 st.error(f"❌ OpenAI API Error: {e}")
#                 return None

# def get_ai_question(resume_text, asked_qs):
#     if not asked_qs:
#         return "Please introduce yourself."
#     prompt = f"""
# You are an AI interviewer. Based on this resume: {resume_text},
# generate **one random interview question** (technical, behavioral, or coding).
# Do NOT repeat previous questions. Previous asked questions: {asked_qs}
# Level: start from basic, then intermediate, then advanced
# """
#     response = call_openai(
#         client.chat.completions.create,
#         model="gpt-4o-mini",
#         messages=[{"role": "user", "content": prompt}]
#     )
#     if response:
#         return response.choices[0].message.content.strip()
#     return "❌ Could not generate question due to API issue."

# def get_feedback(questions_answers):
#     prompt = f"""
# I conducted a mock interview. Here are the questions and answers:
# {questions_answers}

# Give feedback in 2 clear sections:
# ✅ Strengths
# ❌ Improvements
# """
#     response = call_openai(
#         client.chat.completions.create,
#         model="gpt-4o-mini",
#         messages=[{"role": "user", "content": prompt}]
#     )
#     if response:
#         return response.choices[0].message.content.strip()
#     return "❌ Could not generate feedback due to API issue."

# def record_answer(min_seconds=8):
#     recognizer = sr.Recognizer()
#     with sr.Microphone() as source:
#         st.info("🎤 Listening... Please speak now (minimum 8 seconds)")
#         audio = recognizer.listen(source)
#         try:
#             text = recognizer.recognize_google(audio)
#             duration = len(text.split()) / 2
#             if duration < min_seconds:
#                 st.warning(f"⚠️ Your answer is too short. Try at least {min_seconds} seconds.")
#             st.success(f"📝 Recognized: {text}")
#             return text
#         except:
#             st.error("⚠️ Could not recognize your speech. Try again.")
#             return None

# # ===============================
# # STREAMLIT APP
# # ===============================
# st.title("🧠 AI Mock Interviewer with Speech & Text Input")

# # Session state
# if "resume_text" not in st.session_state:
#     st.session_state.resume_text = ""
# if "questions" not in st.session_state:
#     st.session_state.questions = []
# if "answers" not in st.session_state:
#     st.session_state.answers = []
# if "interview_finished" not in st.session_state:
#     st.session_state.interview_finished = False

# # Upload Resume
# uploaded_file = st.file_uploader("📂 Upload your Resume (PDF/DOCX/TXT)", type=["pdf", "docx", "txt"])
# if uploaded_file and not st.session_state.resume_text:
#     st.session_state.resume_text = extract_text(uploaded_file)
#     st.success("✅ Resume uploaded & processed!")

# # Start Interview
# if st.session_state.resume_text and not st.session_state.interview_finished:

#     if st.button("🤖 Ask Next Question"):
#         q = get_ai_question(st.session_state.resume_text, st.session_state.questions)
#         st.session_state.questions.append(q)
#         st.session_state.answers.append("")

#     if st.session_state.questions:
#         q_index = len(st.session_state.questions) - 1
#         st.markdown(f'<div class="qa-card"><h4>Q{q_index+1}: {st.session_state.questions[q_index]}</h4></div>', unsafe_allow_html=True)

#         answer_choice = st.radio("Choose your answer method:", ("Type Answer", "Speak Answer"))

#         if answer_choice == "Type Answer":
#             typed_answer_key = f"typed_answer_{q_index}"
#             if typed_answer_key not in st.session_state:
#                 st.session_state[typed_answer_key] = ""
#             typed_answer = st.text_area(
#                 "📝 Write your answer here:",
#                 value=st.session_state[typed_answer_key],
#                 height=150,
#                 key=typed_answer_key
#             )
#             if st.button("💾 Save Answer"):
#                 st.session_state.answers[q_index] = typed_answer
#                 st.success("✅ Answer saved!")
#                 st.session_state[typed_answer_key] = ""

#         elif answer_choice == "Speak Answer":
#             if st.button("🎤 Record My Answer"):
#                 ans = record_answer()
#                 if ans:
#                     st.session_state.answers[q_index] = ans

#     # Show Q&A so far
#     if st.session_state.answers:
#         st.subheader("📌 Your Progress")
#         for i, (q, a) in enumerate(zip(st.session_state.questions, st.session_state.answers)):
#             st.markdown(f'<div class="qa-card"><b>Q{i+1}:</b> {q}<br><b>A{i+1}:</b> {a if a else "❌ Not answered yet"}</div>', unsafe_allow_html=True)

#     # Stop interview
#     if st.button("🛑 Stop Interview & Get Feedback"):
#         st.session_state.interview_finished = True

# # Finish Interview automatically after 10 questions or stopped manually
# if st.session_state.interview_finished or len(st.session_state.questions) >= 10:
#     qa_text = "\n".join([f"Q: {q}\nA: {a}" for q, a in zip(st.session_state.questions, st.session_state.answers)])
#     feedback = get_feedback(qa_text)
#     st.markdown(f'<div class="feedback-box"><h3>📋 Interview Feedback</h3><p>{feedback}</p></div>', unsafe_allow_html=True)

























# import streamlit as st
# import speech_recognition as sr
# import PyPDF2
# import docx
# import time
# from openai import OpenAI
# from openai import OpenAIError

# # ===============================
# # CONFIG
# # ===============================
# st.set_page_config(page_title="AI Mock Interviewer", page_icon="🧠", layout="wide")

# # ===============================
# # CSS for Clear Buttons & Theme
# # ===============================
# st.markdown("""
# <style>
# /* App background */
# [data-testid="stAppViewContainer"] {
#     background-color:#CB94F7;
#     color: #f0f0f0;
#     font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
# }

# /* Main title */
# h1 {
#     text-align: center;
#     font-size: 3rem;
#     color: #00ffea;
#     margin-bottom: 20px;
# }

# /* Buttons */
# div.stButton > button {
#     # background-color:#FFC0CB;
#     # color: #121212;
#     font-weight: bold;
#     border-radius: 10px;
#     padding: 10px 20px;
#     font-size: 16px;
#     transition: all 0.2s ease-in-out;
#     width: 100%;
# }
# div.stButton > button:hover {
#     background-color:pink;
#     transform: scale(1.05);
# }

# /* Radio buttons */
# div[role="radiogroup"] label {
#     color: #00ffea;
#     font-weight: 600;
# }

# /* Text area */
# textarea {
#     background: #1f1f1f !important;
#     color: #f0f0f0 !important;
#     border-radius: 10px !important;
#     padding: 10px !important;
#     font-size: 15px !important;
# }

# /* File uploader label */
# div.stFileUploader label {
#     color: #00ffea !important;
#     font-weight: 600;
#     font-size: 1.1rem;
# }

# /* Q&A card */
# .qa-card {
#     background-color: #1f1f1f;
#     border-radius: 12px;
#     padding: 15px;
#     margin: 10px 0;
#     color: #f0f0f0;
#     box-shadow: 0px 4px 12px rgba(0,255,234,0.2);
# }

# /* Feedback box */
# .feedback-box {
#     background-color: #1f1f1f;
#     border-radius: 12px;
#     padding: 20px;
#     margin-top: 15px;
#     color: #f0f0f0;
#     box-shadow: 0px 4px 15px rgba(0,255,234,0.3);
# }
# </style>
# """, unsafe_allow_html=True)

# # ===============================
# # OPENAI CONFIG
# # ===============================
# client = OpenAI(api_key="YOUR_OPENAI_API_KEY")  # Replace with your API key

# # ===============================
# # FUNCTIONS
# # ===============================
# def extract_text(file):
#     text = ""
#     if file.type == "application/pdf":
#         reader = PyPDF2.PdfReader(file)
#         for page in reader.pages:
#             text += page.extract_text() or ""
#     elif file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
#         doc = docx.Document(file)
#         for para in doc.paragraphs:
#             text += para.text + "\n"
#     else:
#         text = file.read().decode("utf-8")
#     return text.strip()

# def call_openai(func, *args, retries=3, delay=5, **kwargs):
#     for attempt in range(retries):
#         try:
#             return func(*args, **kwargs)
#         except OpenAIError as e:
#             if "RateLimit" in str(e):
#                 if attempt < retries - 1:
#                     st.warning(f"⚠️ Rate limit reached. Retrying in {delay} seconds...")
#                     time.sleep(delay)
#                 else:
#                     st.error("❌ API rate limit exceeded. Try again later.")
#                     return None
#             else:
#                 st.error(f"❌ OpenAI API Error: {e}")
#                 return None

# def get_ai_question(resume_text, asked_qs):
#     if not asked_qs:
#         return "Please introduce yourself."
#     prompt = f"""
# You are an AI interviewer. Based on this resume: {resume_text},
# generate **one random interview question** (technical, behavioral, or coding).
# Do NOT repeat previous questions. Previous asked questions: {asked_qs}
# Level: start from basic, then intermediate, then advanced
# """
#     response = call_openai(
#         client.chat.completions.create,
#         model="gpt-4o-mini",
#         messages=[{"role": "user", "content": prompt}]
#     )
#     if response:
#         return response.choices[0].message.content.strip()
#     return "❌ Could not generate question due to API issue."

# def get_feedback(questions_answers):
#     prompt = f"""
# I conducted a mock interview. Here are the questions and answers:
# {questions_answers}

# Give feedback in 2 clear sections:
# ✅ Strengths
# ❌ Improvements
# """
#     response = call_openai(
#         client.chat.completions.create,
#         model="gpt-4o-mini",
#         messages=[{"role": "user", "content": prompt}]
#     )
#     if response:
#         return response.choices[0].message.content.strip()
#     return "❌ Could not generate feedback due to API issue."

# def record_answer(min_seconds=8):
#     recognizer = sr.Recognizer()
#     with sr.Microphone() as source:
#         st.info("🎤 Listening... Please speak now (minimum 8 seconds)")
#         audio = recognizer.listen(source)
#         try:
#             text = recognizer.recognize_google(audio)
#             duration = len(text.split()) / 2
#             if duration < min_seconds:
#                 st.warning(f"⚠️ Your answer is too short. Try at least {min_seconds} seconds.")
#             st.success(f"📝 Recognized: {text}")
#             return text
#         except:
#             st.error("⚠️ Could not recognize your speech. Try again.")
#             return None

# # ===============================
# # STREAMLIT APP
# # ===============================
# st.title("🧠 AI Mock Interviewer")

# # Session state
# if "resume_text" not in st.session_state:
#     st.session_state.resume_text = ""
# if "questions" not in st.session_state:
#     st.session_state.questions = []
# if "answers" not in st.session_state:
#     st.session_state.answers = []
# if "interview_finished" not in st.session_state:
#     st.session_state.interview_finished = False

# # Upload Resume
# uploaded_file = st.file_uploader("📂 Upload your Resume (PDF/DOCX/TXT)", type=["pdf", "docx", "txt"])
# if uploaded_file and not st.session_state.resume_text:
#     st.session_state.resume_text = extract_text(uploaded_file)
#     st.success("✅ Resume uploaded & processed!")

# # Start Interview
# if st.session_state.resume_text and not st.session_state.interview_finished:

#     if st.button("🤖 Ask Next Question"):
#         q = get_ai_question(st.session_state.resume_text, st.session_state.questions)
#         st.session_state.questions.append(q)
#         st.session_state.answers.append("")

#     if st.session_state.questions:
#         q_index = len(st.session_state.questions) - 1
#         st.markdown(f'<div class="qa-card"><h4>Q{q_index+1}: {st.session_state.questions[q_index]}</h4></div>', unsafe_allow_html=True)

#         answer_choice = st.radio("Choose your answer method:", ("Type Answer", "Speak Answer"))

#         if answer_choice == "Type Answer":
#             typed_answer_key = f"typed_answer_{q_index}"
#             if typed_answer_key not in st.session_state:
#                 st.session_state[typed_answer_key] = ""
#             typed_answer = st.text_area(
#                 "📝 Write your answer here:",
#                 value=st.session_state[typed_answer_key],
#                 height=150,
#                 key=typed_answer_key
#             )
#             if st.button("💾 Save Answer"):
#                 st.session_state.answers[q_index] = typed_answer
#                 st.success("✅ Answer saved!")
#                 st.session_state[typed_answer_key] = ""

#         elif answer_choice == "Speak Answer":
#             if st.button("🎤 Record My Answer"):
#                 ans = record_answer()
#                 if ans:
#                     st.session_state.answers[q_index] = ans

#     # Show Q&A so far
#     if st.session_state.answers:
#         st.subheader("📌 Your Progress")
#         for i, (q, a) in enumerate(zip(st.session_state.questions, st.session_state.answers)):
#             st.markdown(f'<div class="qa-card"><b>Q{i+1}:</b> {q}<br><b>A{i+1}:</b> {a if a else "❌ Not answered yet"}</div>', unsafe_allow_html=True)

#     # Stop interview
#     if st.button("🛑 Stop Interview & Get Feedback"):
#         st.session_state.interview_finished = True

# # Finish Interview automatically after 10 questions or stopped manually
# if st.session_state.interview_finished or len(st.session_state.questions) >= 10:
#     qa_text = "\n".join([f"Q: {q}\nA: {a}" for q, a in zip(st.session_state.questions, st.session_state.answers)])
#     feedback = get_feedback(qa_text)
#     st.markdown(f'<div class="feedback-box"><h3>📋 Interview Feedback</h3><p>{feedback}</p></div>', unsafe_allow_html=True)





















import streamlit as st
import speech_recognition as sr
import PyPDF2
import docx
import time
from openai import OpenAI
from openai import OpenAIError

# ===============================
# CONFIG
# ===============================
st.set_page_config(page_title="AI Mock Interviewer", page_icon="🧠", layout="wide")

# ===============================
# CSS for Navbar & Theme
# ===============================
st.markdown("""

<style>
/* Background with image */
[data-testid="stAppViewContainer"] {
    background: url('https://images.unsplash.com/photo-1522202176988-66273c2fd55f') no-repeat center center fixed;
    background-size: cover;
    color: #FFD700; /* Yellow text */
    font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
}

/* Navbar */
.navbar {
    width: 100%;
    background-color: #87CEFA; /* Light Blue */
    padding: 12px 20px;
    position: sticky;
    top: 0;
    z-index: 999;
    display: flex;
    justify-content: space-between;
    align-items: center;
    box-shadow: 0px 2px 15px rgba(0,0,0,0.3);
    border-bottom: 2px solid #FFD700; /* Yellow border */
    border-radius: 0 0 15px 15px;
}
.navbar h2 {
    color: #FFD700; /* Yellow */
    margin: 0;
    font-size: 1.8rem;
    font-weight: bold;
}
.navbar button {
    background-color: #FFD700; /* Yellow */
    color: #000000;
    font-weight: bold;
    border-radius: 10px;
    padding: 8px 18px;
    font-size: 15px;
    border: 1px solid #87CEFA; /* Light Blue border */
    cursor: pointer;
    transition: all 0.3s ease-in-out;
}
.navbar button:hover {
    background-color: #f5c400; /* Darker Yellow */
    transform: scale(1.05);
}

/* Main title */
h1 {
    text-align: center;
    font-size: 3rem;
    color: #FFD700; /* Yellow */
    margin-bottom: 20px;
    text-shadow: 2px 2px 8px rgba(0,0,0,0.6);
}

/* Buttons */
div.stButton > button {
    background-color: #87CEFA; /* Light Blue */
    color: #000000 !important; /* Black text for contrast */
    font-weight: bold;
    border-radius: 12px;
    padding: 10px 22px;
    font-size: 17px;
    transition: all 0.3s ease-in-out;
    width: 100%;
    border: 2px solid #FFD700 !important; /* Yellow border */
}
div.stButton > button:hover {
    background-color: #5db7e6 !important; /* Slightly darker blue */
    transform: scale(1.07);
}
/* Login Form */
.login-box {
    background-color: rgba(135, 206, 250, 0.85); /* Semi-transparent light blue */
    border-radius: 15px;
    padding: 20px; /* Smaller padding */
    max-width: 220px; /* Smaller box size */
    margin: auto;
    margin-top: 60px;
    box-shadow: 0px 6px 20px rgba(0,0,0,0.4);
    border: 2px solid #FFD700; /* Yellow border */
    text-align: center;
    transform: scale(0.9); /* Slightly shrink the whole box */
}

/* Login Heading */
.login-box h2 {
    font-size: 2.5rem; /* Bigger heading */
    font-weight: 900; /* Extra bold */
    color: #ff4500; /* Bright orange-red for contrast */
    margin-bottom: 20px;
    text-shadow: 2px 2px 12px rgba(0,0,0,0.7);
    letter-spacing: 1px;
}

/* Labels */
.login-box label {
    font-size: 1.2rem;
    font-weight: bold;
    color: #FFD700; /* Yellow */
    display: block;
    margin: 8px 0 4px;
}

/* Inputs */
.login-box input {
    width: 70%; /* Smaller input width */
    padding: 6px; /* Smaller padding */
    font-size: 0.9rem; /* Slightly smaller text */
    border-radius: 8px;
    border: 1px solid #FFD700; /* Yellow border */
    background-color: #87CEFA; /* Light Blue */
    color: #000000; /* Black text */
    margin-bottom: 12px;
}
.login-box input:focus {
    outline: none;
    border: 2px solid #FFD700;
    box-shadow: 0px 0px 10px #FFD700;
}

</style>


""", unsafe_allow_html=True)

# ===============================
# OPENAI CONFIG
# ===============================
client = OpenAI(api_key="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")  # Replace with your API key

# ===============================
# FUNCTIONS
# ===============================
def extract_text(file):
    text = ""
    if file.type == "application/pdf":
        reader = PyPDF2.PdfReader(file)
        for page in reader.pages:
            text += page.extract_text() or ""
    elif file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        doc = docx.Document(file)
        for para in doc.paragraphs:
            text += para.text + "\n"
    else:
        text = file.read().decode("utf-8")
    return text.strip()

def call_openai(func, *args, retries=3, delay=5, **kwargs):
    for attempt in range(retries):
        try:
            return func(*args, **kwargs)
        except OpenAIError as e:
            if "RateLimit" in str(e):
                if attempt < retries - 1:
                    st.warning(f"⚠️ Rate limit reached. Retrying in {delay} seconds...")
                    time.sleep(delay)
                else:
                    st.error("❌ API rate limit exceeded. Try again later.")
                    return None
            else:
                st.error(f"❌ OpenAI API Error: {e}")
                return None

def get_ai_question(resume_text, asked_qs):
    if not asked_qs:
        return "Please introduce yourself."
    prompt = f"""
You are an AI interviewer. Based on this resume: {resume_text},
generate **one random interview question** (technical, behavioral, or coding).
Do NOT repeat previous questions. Previous asked questions: {asked_qs}
Level: start from basic, then intermediate, then advanced
"""
    response = call_openai(
        client.chat.completions.create,
        model="gpt-4o-mini",
        messages=[{"role": "user", "content": prompt}]
    )
    if response:
        return response.choices[0].message.content.strip()
    return "❌ Could not generate question due to API issue."

def get_feedback(questions_answers):
    prompt = f"""
I conducted a mock interview. Here are the questions and answers:
{questions_answers}

Give feedback in 2 clear sections:
✅ Strengths
❌ Improvements
"""
    response = call_openai(
        client.chat.completions.create,
        model="gpt-4o-mini",
        messages=[{"role": "user", "content": prompt}]
    )
    if response:
        return response.choices[0].message.content.strip()
    return "❌ Could not generate feedback due to API issue."

def record_answer(min_seconds=8):
    recognizer = sr.Recognizer()
    with sr.Microphone() as source:
        st.info("🎤 Listening... Please speak now (minimum 8 seconds)")
        audio = recognizer.listen(source)
        try:
            text = recognizer.recognize_google(audio)
            duration = len(text.split()) / 2
            if duration < min_seconds:
                st.warning(f"⚠️ Your answer is too short. Try at least {min_seconds} seconds.")
            st.success(f"📝 Recognized: {text}")
            return text
        except:
            st.error("⚠️ Could not recognize your speech. Try again.")
            return None

# ===============================
# SESSION STATE SETUP
# ===============================
if "logged_in" not in st.session_state:
    st.session_state.logged_in = False
if "user_info" not in st.session_state:
    st.session_state.user_info = {}
if "resume_text" not in st.session_state:
    st.session_state.resume_text = ""
if "questions" not in st.session_state:
    st.session_state.questions = []
if "answers" not in st.session_state:
    st.session_state.answers = []
if "interview_finished" not in st.session_state:
    st.session_state.interview_finished = False

# ===============================
# LOGIN PAGE
# ===============================
if not st.session_state.logged_in:
    st.markdown('<h1>🧠 AI Mock Interviewer Login</h1>', unsafe_allow_html=True)
    with st.form("login_form"):
        name = st.text_input("Full Name")
        course = st.text_input("Course Appearing")
        branch = st.text_input("Branch")
        location = st.text_input("Your Location")
        login_btn = st.form_submit_button("Login")
    if login_btn:
        if name and course and branch and location:
            st.session_state.logged_in = True
            st.session_state.user_info = {
                "name": name,
                "course": course,
                "branch": branch,
                "location": location
            }
            st.success(f"Welcome {name}! You can now proceed to the interview.")
        else:
            st.warning("Please fill all the fields to login.")

# ===============================
# MAIN APP PAGE
# ===============================
if st.session_state.logged_in:

    # Navbar
    st.markdown(f'''
    <div class="navbar">
        <h2>AI Mock Interviewer</h2>
        <div>
            <span>Welcome, {st.session_state.user_info['name']} ({st.session_state.user_info['course']})</span>
            <button onclick="window.location.reload();">Logout</button>
        </div>
    </div>
    ''', unsafe_allow_html=True)

    st.title("📝 Upload Resume & Start Interview")

    # ===============================
    # Resume Upload & Interview Logic
    # ===============================
    uploaded_file = st.file_uploader("📂 Upload your Resume (PDF/DOCX/TXT)", type=["pdf", "docx", "txt"])
    if uploaded_file and not st.session_state.resume_text:
        st.session_state.resume_text = extract_text(uploaded_file)
        st.success("✅ Resume uploaded & processed!")

    # Start Interview
    if st.session_state.resume_text and not st.session_state.interview_finished:

        if st.button("🤖 Ask Next Question"):
            q = get_ai_question(st.session_state.resume_text, st.session_state.questions)
            st.session_state.questions.append(q)
            st.session_state.answers.append("")

        if st.session_state.questions:
            q_index = len(st.session_state.questions) - 1
            st.markdown(f'<div class="qa-card"><h4>Q{q_index+1}: {st.session_state.questions[q_index]}</h4></div>', unsafe_allow_html=True)

            answer_choice = st.radio("Choose your answer method:", ("Type Answer", "Speak Answer"))

            if answer_choice == "Type Answer":
                typed_answer_key = f"typed_answer_{q_index}"
                if typed_answer_key not in st.session_state:
                    st.session_state[typed_answer_key] = ""
                typed_answer = st.text_area(
                    "📝 Write your answer here:",
                    value=st.session_state[typed_answer_key],
                    height=150,
                    key=typed_answer_key
                )
                if st.button("💾 Save Answer"):
                    st.session_state.answers[q_index] = typed_answer
                    st.success("✅ Answer saved!")
                    st.session_state[typed_answer_key] = ""

            elif answer_choice == "Speak Answer":
                if st.button("🎤 Record My Answer"):
                    ans = record_answer()
                    if ans:
                        st.session_state.answers[q_index] = ans

        # Show Q&A so far
        if st.session_state.answers:
            st.subheader("📌 Your Progress")
            for i, (q, a) in enumerate(zip(st.session_state.questions, st.session_state.answers)):
                st.markdown(f'<div class="qa-card"><b>Q{i+1}:</b> {q}<br><b>A{i+1}:</b> {a if a else "❌ Not answered yet"}</div>', unsafe_allow_html=True)






























